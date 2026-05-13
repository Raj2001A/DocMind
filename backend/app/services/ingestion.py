"""
ingestion.py
------------
Document ingestion pipeline — Phase 1 & 2.

Pipeline:
  PDF / Docx  ──► text extraction
                ──► smart chunking (512 tokens / 50 overlap)
                ──► OpenAI embeddings
                ──► ChromaDB storage
                ──► BM25 index update (in-memory, rebuilt per session)

Interview talking point:
  "I use RecursiveCharacterTextSplitter instead of fixed-size splitting
   because it respects sentence and paragraph boundaries, which improves
   retrieval coherence. The 50-token overlap ensures no context is lost
   at chunk boundaries."
"""
import hashlib
import logging
import pickle
import uuid
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from rank_bm25 import BM25Okapi

from app.config import settings
from app.models.db import save_document

logger = logging.getLogger(__name__)

# ── BM25 index (in-memory, persisted to disk as pickle) ──────────────────────
BM25_INDEX_PATH = Path("./bm25_index.pkl")
_bm25_corpus: list[list[str]] = []
_bm25_doc_map: list[dict] = []   # parallel list: metadata for each BM25 doc
_bm25_index: Optional[BM25Okapi] = None

# ── ChromaDB Singleton ────────────────────────────────────────────────────────
# Lazily initialise once and reuse — safe under uvicorn single-process.
_vectorstore: Optional["Chroma"] = None


def _get_embeddings():
    """
    Provider-agnostic embeddings factory.
    Mirrors the LLM factory pattern in agents/graph.py — same env var controls both.
    """
    provider = settings.llm_provider.lower()
    if provider == "gemini":
        from langchain_google_genai import GoogleGenerativeAIEmbeddings
        return GoogleGenerativeAIEmbeddings(
            model=settings.gemini_embedding_model,
            google_api_key=settings.gemini_api_key,
        )
    else:  # openai (default)
        return OpenAIEmbeddings(
            model=settings.openai_embedding_model,
            openai_api_key=settings.openai_api_key,
        )


def get_vectorstore() -> "Chroma":
    """Return the singleton ChromaDB vectorstore, initialising it on first call."""
    global _vectorstore
    if _vectorstore is None:
        embeddings = _get_embeddings()
        _vectorstore = Chroma(
            collection_name="docmind",
            embedding_function=embeddings,
            persist_directory=settings.chroma_persist_dir,
        )
        logger.info(f"ChromaDB vectorstore initialised (provider={settings.llm_provider})")
    return _vectorstore


def _load_bm25() -> None:
    global _bm25_corpus, _bm25_doc_map, _bm25_index
    if BM25_INDEX_PATH.exists():
        with open(BM25_INDEX_PATH, "rb") as f:
            data = pickle.load(f)
        _bm25_corpus = data["corpus"]
        _bm25_doc_map = data["doc_map"]
        _bm25_index = BM25Okapi(_bm25_corpus) if _bm25_corpus else None
        logger.info(f"BM25: loaded {len(_bm25_corpus)} documents from disk")


def _save_bm25() -> None:
    with open(BM25_INDEX_PATH, "wb") as f:
        pickle.dump({"corpus": _bm25_corpus, "doc_map": _bm25_doc_map}, f)


def _update_bm25(chunks: list[Document]) -> None:
    global _bm25_corpus, _bm25_doc_map, _bm25_index
    for chunk in chunks:
        tokens = chunk.page_content.lower().split()
        _bm25_corpus.append(tokens)
        # Store chunk_text alongside metadata so retrieval.py can
        # reconstruct full Document objects from BM25 results.
        meta = {**chunk.metadata, "chunk_text": chunk.page_content}
        _bm25_doc_map.append(meta)
    _bm25_index = BM25Okapi(_bm25_corpus)
    _save_bm25()


# ── Text extraction ───────────────────────────────────────────────────────────
def _extract_pdf(path: Path) -> list[tuple[str, int]]:
    """Returns list of (text, page_number) tuples."""
    pages = []
    doc = fitz.open(str(path))
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            pages.append((text, i + 1))
    return pages


def _extract_docx(path: Path) -> list[tuple[str, int]]:
    doc = docx.Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [(full_text, 1)]  # Word docs treated as single page


# ── Chunking ──────────────────────────────────────────────────────────────────
def _chunk_pages(
    pages: list[tuple[str, int]],
    filename: str,
    document_id: str,
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks: list[Document] = []
    for page_text, page_num in pages:
        sub_chunks = splitter.split_text(page_text)
        for i, chunk_text in enumerate(sub_chunks):
            chunks.append(
                Document(
                    page_content=chunk_text,
                    metadata={
                        "document_id": document_id,
                        "filename": filename,
                        "page": page_num,
                        "chunk_index": len(chunks) + i,
                    },
                )
            )
    return chunks


# ── Vectorstore (delegates to singleton in ingestion.py) ─────────────────────
def _get_vectorstore() -> Chroma:
    from app.services.ingestion import get_vectorstore
    return get_vectorstore()


# ── Public API ────────────────────────────────────────────────────────────────
def ingest_document(file_path: Path, filename: str) -> dict:
    """
    Full ingestion pipeline. Returns metadata dict including chunk_count.
    """
    _load_bm25()

    document_id = hashlib.md5(filename.encode()).hexdigest()[:12] + str(uuid.uuid4())[:8]

    # 1. Extract
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        pages = _extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        pages = _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    logger.info(f"Extracted {len(pages)} pages from '{filename}'")

    # 2. Chunk
    chunks = _chunk_pages(pages, filename, document_id)
    logger.info(f"Created {len(chunks)} chunks")

    # 3. Store in ChromaDB (uses singleton — no reconnect cost)
    vectorstore = get_vectorstore()
    vectorstore.add_documents(chunks)
    logger.info(f"Stored {len(chunks)} chunks in ChromaDB")

    # 4. Update BM25 index
    _update_bm25(chunks)
    logger.info(f"BM25 index updated — total docs: {len(_bm25_corpus)}")

    # 5. Save to SQLite
    save_document(document_id, filename, len(chunks))

    return {
        "document_id": document_id,
        "filename": filename,
        "chunk_count": len(chunks),
    }


def get_bm25_state() -> tuple[Optional[BM25Okapi], list[dict]]:
    """Returns the current BM25 index and its document map."""
    _load_bm25()
    return _bm25_index, _bm25_doc_map
